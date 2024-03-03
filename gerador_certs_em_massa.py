# -*- coding: utf-8 -*-
import json
from model.Certificado import Certificado
from utils.DesenhaCertificado import DesenhaCertificado
from docx import Document


def config_cert() -> list:
    with open('./utils/config_cert.json', encoding='utf-8') as config_json:
        config_certificado = json.load(config_json)
        return config_certificado


def certs() -> list:
    with open('./utils/cert.json', encoding='utf-8') as config_json:
        config_certificado = json.load(config_json)
        return config_certificado


def process_certificado_step():
    users = certs()
    c = 0
    for obj in users:
        cert = Certificado(
            obj['empresa'],
            obj['nomeDivulgacao'],
            obj['segmento'],
            obj['cidade'],
            obj['uf'],
            obj['retroativos']
        )
        draw = DesenhaCertificado(cert, config_cert())
        draw.criar_certificado(modo_bot=False)
        if len(cert.retroativos) > 0:
            for certObj in cert.retroativos:
                cert.ano = certObj
                draw = DesenhaCertificado(cert, config_cert())
                draw.criar_certificado(modo_bot=False)

        c += 1
        print(c)


def process_lista_entrega(users):
    texts = []
    for obj in users:
        segmento = obj.get('segmento', '')
        empresa = obj.get('nomeDivulgacao', '') or obj.get('empresa', '')
        preco = obj.get('preco', '')
        obs = obj.get('obs', '')

        text = f'{segmento} = {empresa} = {preco} = {obs}'
        texts.append(text)

    cidade = users[0].get('cidade', 'desconhecida') if users else 'desconhecida'
    filename = f'{cidade}-lista-entrega.docx'
    return texts, filename


def process_lista_divulgacao(users):
    texts = []
    for obj in users:
        segmento = obj.get('segmento', '')
        empresa = obj.get('nomeDivulgacao', '') or obj.get('empresa', '')

        text = f'{segmento} = {empresa}'
        texts.append(text)

    cidade = users[0].get('cidade', 'desconhecida') if users else 'desconhecida'
    filename = f'{cidade}-lista-divulgação.docx'
    return texts, filename


# Exemplo de uso:
users = certs()
texts_entrega, filename_entrega = process_lista_entrega(users)
texts_divulgacao, filename_divulgacao = process_lista_divulgacao(users)


def save_to_docx(texts, filename):
    doc = Document()
    for text in texts:
        paragrafo = doc.add_paragraph()
        paragrafo.add_run(text).bold = True
    doc.save(filename)

if __name__ == '__main__':
    process_certificado_step()
    save_to_docx(texts_entrega, filename_entrega)
    save_to_docx(texts_divulgacao, filename_divulgacao)