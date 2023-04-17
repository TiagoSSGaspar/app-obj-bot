import json
from model.Certificado import Certificado
from utils.DesenhaCertificado import DesenhaCertificado
from docx import Document


def config_cert() -> list:
    with open('./utils/config_cert.json') as config_json:
        config_certificado = json.load(config_json)
        return config_certificado


def certs() -> list:
    with open('./utils/cert.json') as config_json:
        config_certificado = json.load(config_json)
        return config_certificado


def process_certificado_step():
    users = certs()
    c = 0
    for obj in users:
        cert = Certificado(
            obj['empresa'],
            obj['nomeDivulgacao'],
            obj['segmento'],
            obj['cidade'],
            obj['uf'],
            obj['retroativos']
        )
        draw = DesenhaCertificado(cert, config_cert())
        draw.criar_certificado()
        if len(cert.retroativos) > 0:
            for certObj in cert.retroativos:
                cert.ano = certObj
                draw = DesenhaCertificado(cert, config_cert())
                draw.criar_certificado()

        c += 1
        print(c)


documento = Document()



def process_lista_entrega():
    users = certs()
    c = 0
    for obj in users:
        empresa = obj['empresa']
        segmento = obj['segmento']
        nome_divulgacao = obj['nomeDivulgacao']
        preco = obj['preco']
        obs = obj['obs']
        if nome_divulgacao != '':
            empresa = nome_divulgacao

        text = '{} = {}'.format(segmento, empresa)
        #text = '{} = {} = {} = {}'.format(segmento, empresa, preco, obs)

        paragrafo = documento.add_paragraph()
        paragrafo.add_run(text).bold = True

    #documento.save('lista-entrega.docx')
    documento.save('lista-divulgação.docx')


#process_certificado_step()
process_lista_entrega()
