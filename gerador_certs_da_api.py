## -*- coding: utf-8 -*-
import json
import requests
import sys
import io

from model.Certificado import Certificado
from utils.DesenhaCertificado import DesenhaCertificado
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')


def config_cert() -> list:
    with open('./utils/config_cert.json', encoding='utf-8') as config_json:
        config_certificado = json.load(config_json)
        return config_certificado


def fetch_and_select_item(url: str):
    try:
        # Buscar dados da API
        response = requests.get(url)
        response.raise_for_status()
        data = response.json()

        # Verificar se a resposta contém uma lista de dicionários
        if not isinstance(data, list):
            print("Dados não são uma lista.")
            return

        if not data:
            print("Nenhuma venda encontrada.")
            return

        # Exibir lista de vendas
        for idx, item in enumerate(data):
            # Exibir uma chave específica do dicionário, ajuste conforme necessário
            print(f"{idx + 1}: {item.get('cityName', 'N/A')}")

        # Obter seleção do usuário
        while True:
            try:
                choice = int(input(f"Escolha um item (1-{len(data)}): "))
                if 1 <= choice <= len(data):
                    selected_item = data[choice - 1]
                    print(f"Você selecionou: {selected_item}")
                    return selected_item
                else:
                    print("Escolha inválida, tente novamente.")
            except ValueError:
                print("Entrada inválida, por favor insira um número.")

    except requests.RequestException as e:
        print(f"Erro ao buscar dados da API: {e}")


def certs(obj_sales) -> list:
    list_certs = []
    sales = obj_sales['sales']

    for obj in sales:
        cert = Certificado(
            nome_fantasia=obj['company'],
            nome_divulgacao=obj['divulgationName'],
            segmento=obj['segment'],
            cidade=obj_sales['cityName'],
            uf=obj_sales['uf'],
            retroativos=obj['retroactiveCertificates'],
            preco=obj['amount'],
            obs=obj['obs'],
        )

        list_certs.append(cert)

    return list_certs


obj_sales = fetch_and_select_item('https://api-obj-for-android.onrender.com/sales-register')
users = certs(obj_sales)

def process_certificado_step():
    c = 0
    for cert in users:
        draw = DesenhaCertificado(cert, config_cert())
        draw.criar_certificado(modo_bot=False)
        if len(cert.retroativos) > 0:
            for certObj in cert.retroativos:
                cert.ano = certObj
                draw = DesenhaCertificado(cert, config_cert())
                draw.criar_certificado(modo_bot=False)

        c += 1
        print(c)


def process_lista_entrega(users):
    texts = []
    for obj in users:
        segmento = obj.segmento
        empresa = obj.nome_fantasia
        preco = obj.preco
        obs = obj.obs

        text = f'{segmento} = {empresa} = {preco} = {obs}'
        texts.append(text)

    cidade = users[0].cidade if users else 'desconhecida'
    filename = f'{cidade}-lista-entrega.docx'
    return texts, filename


def process_lista_divulgacao(users):
    texts = []
    for obj in users:
        segmento = obj.segmento
        empresa = obj.nome_fantasia or obj.nome_divulgacao

        text = f'{segmento} = {empresa}'
        texts.append(text)

    cidade = users[0].cidade if users else 'desconhecida'
    filename = f'{cidade}-lista-divulgação.docx'
    return texts, filename


# Exemplo de uso:
texts_entrega, filename_entrega = process_lista_entrega(users)
texts_divulgacao, filename_divulgacao = process_lista_divulgacao(users)


def save_to_docx(texts, filename):
    doc = Document()
    for text in texts:
        paragrafo = doc.add_paragraph()
        paragrafo.add_run(text).bold = True
    doc.save(filename)


if __name__ == '__main__':
    process_certificado_step()
    save_to_docx(texts_entrega, filename_entrega)
    save_to_docx(texts_divulgacao, filename_divulgacao)
